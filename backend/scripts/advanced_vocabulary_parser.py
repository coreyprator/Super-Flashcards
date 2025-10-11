"""
Advanced Document Parser for Google Docs vocabulary exports
Extracts structured vocabulary data from .docx files while preserving formatting and images
"""

import re
from pathlib import Path
import json
import csv
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
import argparse
from docx import Document
from docx.shared import Inches
import base64
import io

@dataclass
class VocabularyEntry:
    """Structured vocabulary entry with all extracted information"""
    word_or_phrase: str
    definition: str
    language: str = "french"  # default
    pronunciation: Optional[str] = None
    etymology: Optional[str] = None
    memory_hint: Optional[str] = None
    difficulty_level: int = 1
    related_words: Optional[List[str]] = None
    english_equivalents: Optional[List[str]] = None
    expressions: Optional[List[str]] = None
    context: Optional[str] = None
    image_description: Optional[str] = None
    confidence_score: float = 1.0  # How confident we are in the extraction

class AdvancedVocabularyParser:
    """
    Advanced parser for structured vocabulary documents
    Handles complex French vocabulary with etymology, cognates, and expressions
    """
    
    def __init__(self):
        self.entries = []
        self.current_entry = None
        self.language_map = {
            'français': 'french',
            'french': 'french',
            'anglais': 'english', 
            'english': 'english',
            'espagnol': 'spanish',
            'spanish': 'spanish',
            'allemand': 'german',
            'german': 'german',
            'italien': 'italian',
            'italian': 'italian'
        }
        
    def extract_main_term(self, text: str) -> Optional[str]:
        """Extract the main vocabulary term from text"""
        # Look for terms in quotes or at start of paragraph
        if '"' in text:
            match = re.search(r'"([^"]+)"', text)
            if match:
                return match.group(1).strip()
        
        # Look for standalone terms (often the first meaningful word)
        words = text.split()
        if words:
            # Clean the first word of punctuation and quotes
            term = re.sub(r'[«»""\(\)\[\]{}]', '', words[0])
            if len(term) > 2 and not term.isupper():
                return term.lower()
        
        return None
    
    def extract_etymology(self, text: str) -> Optional[str]:
        """Extract etymology information"""
        etymology_patterns = [
            r'Étymologie\s*:?\s*(.+?)(?=\n\n|\n[A-Z]|$)',
            r'Origine\s*:?\s*(.+?)(?=\n\n|\n[A-Z]|$)',
            r'vient du latin\s+(.+?)(?=\.|,|\n)',
            r'dérive de\s+(.+?)(?=\.|,|\n)'
        ]
        
        for pattern in etymology_patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                etymology = match.group(1).strip()
                # Clean up bullet points and formatting
                etymology = re.sub(r'●\s*', '• ', etymology)
                etymology = re.sub(r'\s+', ' ', etymology)
                return etymology[:500]  # Limit length
        
        return None
    
    def extract_english_equivalents(self, text: str) -> List[str]:
        """Extract English equivalent words"""
        equivalents = []
        
        # Look for English equivalents section
        patterns = [
            r'Équivalent[^:]*:?\s*["\"]([^"\"]+)["\"]',
            r'en anglais[^:]*:?\s*["\"]([^"\"]+)["\"]',
            r'English[^:]*:?\s*["\"]([^"\"]+)["\"]',
            r'Équivalent\s*:?\s*(.+?)(?=\n●|\n\n|\.$)',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                equiv_text = match.group(1).strip()
                # Split on common separators
                words = re.split(r'[,;]|ou\s+|\s+or\s+', equiv_text)
                for word in words:
                    clean_word = re.sub(r'["\(\)\.]+', '', word.strip())
                    if clean_word and len(clean_word) > 1:
                        equivalents.append(clean_word)
        
        return equivalents[:5]  # Limit to 5 equivalents
    
    def extract_related_words(self, text: str) -> List[str]:
        """Extract French related words"""
        related = []
        
        # Look for related words sections
        patterns = [
            r'Mots apparentés[^:]*:?\s*(.+?)(?=\n\n|\n[A-Z]|$)',
            r'apparentés en français[^:]*:?\s*(.+?)(?=\n\n|\n[A-Z]|$)',
            r'Synonymes?[^:]*:?\s*(.+?)(?=\n\n|\n[A-Z]|$)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                related_text = match.group(1)
                # Extract words from bullet points
                words = re.findall(r'●\s*([^●\n]+)', related_text)
                for word in words:
                    clean_word = re.sub(r'[:\(\)\.]+', '', word.strip())
                    if clean_word and len(clean_word) > 1:
                        related.append(clean_word)
        
        return related[:8]  # Limit to 8 related words
    
    def extract_expressions(self, text: str) -> List[str]:
        """Extract expressions and usage examples"""
        expressions = []
        
        # Look for expressions in quotes
        expr_matches = re.findall(r'["\"]([^"\"]+)["\"]', text)
        for expr in expr_matches:
            if len(expr) > 5 and len(expr) < 100:  # Reasonable expression length
                expressions.append(expr.strip())
        
        # Look for "Expression" sections
        expr_pattern = r'Expression[^:]*:?\s*(.+?)(?=\n\n|\n[A-Z]|$)'
        match = re.search(expr_pattern, text, re.DOTALL | re.IGNORECASE)
        if match:
            expr_text = match.group(1)
            # Extract numbered expressions
            numbered_exprs = re.findall(r'\d+\.\s*["\"]([^"\"]+)["\"]', expr_text)
            expressions.extend(numbered_exprs)
        
        return expressions[:3]  # Limit to 3 expressions
    
    def determine_difficulty(self, entry: VocabularyEntry) -> int:
        """Estimate difficulty level based on content complexity"""
        score = 1
        
        # Complex etymology increases difficulty
        if entry.etymology and len(entry.etymology) > 100:
            score += 1
            
        # Many related words suggests advanced vocabulary
        if entry.related_words and len(entry.related_words) > 4:
            score += 1
            
        # Long definitions suggest complex concepts
        if len(entry.definition) > 200:
            score += 1
            
        # Multiple expressions suggest idiomatic usage
        if entry.expressions and len(entry.expressions) > 1:
            score += 1
        
        return min(score, 5)  # Cap at 5
    
    def parse_vocabulary_section(self, text: str) -> Optional[VocabularyEntry]:
        """Parse a single vocabulary section"""
        lines = text.split('\n')
        if not lines:
            return None
        
        # Try to extract main term from first meaningful line
        main_term = None
        definition_start = 0
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Check if this line contains a standalone term
            if not main_term:
                potential_term = self.extract_main_term(line)
                if potential_term:
                    main_term = potential_term
                    definition_start = i + 1
                    break
        
        if not main_term:
            return None
        
        # Extract definition (text before etymology)
        definition_lines = []
        for i in range(definition_start, len(lines)):
            line = lines[i].strip()
            if any(keyword in line.lower() for keyword in ['étymologie', 'équivalent', 'mots apparentés']):
                break
            if line:
                definition_lines.append(line)
        
        definition = ' '.join(definition_lines).strip()
        if not definition:
            definition = f"Terme français: {main_term}"
        
        # Create entry
        entry = VocabularyEntry(
            word_or_phrase=main_term,
            definition=definition[:500],  # Limit definition length
            language="french"
        )
        
        # Extract additional information
        full_text = text
        entry.etymology = self.extract_etymology(full_text)
        entry.english_equivalents = self.extract_english_equivalents(full_text)
        entry.related_words = self.extract_related_words(full_text)
        entry.expressions = self.extract_expressions(full_text)
        entry.difficulty_level = self.determine_difficulty(entry)
        
        # Calculate confidence based on extracted information
        confidence = 0.5  # Base confidence
        if entry.etymology:
            confidence += 0.2
        if entry.english_equivalents:
            confidence += 0.1
        if entry.related_words:
            confidence += 0.1
        if len(entry.definition) > 50:
            confidence += 0.1
        
        entry.confidence_score = min(confidence, 1.0)
        
        return entry
    
    def split_document_sections(self, text: str) -> List[str]:
        """Split document into vocabulary sections"""
        # Split on major section breaks
        sections = []
        
        # Look for patterns that indicate new vocabulary entries
        # This is based on the structure you showed
        patterns = [
            r'\n\s*\n(?=[A-Za-zÀ-ÿ][a-zà-ÿ]*\s+["\"])',  # New term in quotes
            r'\n\s*\n(?=\d+\.\d+\.?\s+[a-zà-ÿ]+)',  # Numbered sections like "17.1. chaleureuse"
            r'\n\s*\n(?=[A-Za-zÀ-ÿ][a-zà-ÿ]+\s*\n)',  # Standalone terms
            r'_{8,}',  # Underline separators
        ]
        
        # Try each pattern
        for pattern in patterns:
            potential_sections = re.split(pattern, text)
            if len(potential_sections) > len(sections):
                sections = potential_sections
        
        # If no clear pattern, split on double newlines
        if len(sections) <= 1:
            sections = re.split(r'\n\s*\n\s*\n', text)
        
        # Clean and filter sections
        clean_sections = []
        for section in sections:
            section = section.strip()
            if len(section) > 50:  # Minimum section length
                clean_sections.append(section)
        
        return clean_sections
    
    def parse_docx(self, file_path: Path) -> List[VocabularyEntry]:
        """Parse a .docx file and extract vocabulary entries"""
        print(f"📖 Parsing document: {file_path}")
        
        doc = Document(file_path)
        
        # Extract all text
        all_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)
        
        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        all_text.append(cell.text)
        
        full_text = '\n'.join(all_text)
        print(f"📄 Extracted {len(full_text)} characters of text")
        
        # Split into vocabulary sections
        sections = self.split_document_sections(full_text)
        print(f"📋 Found {len(sections)} potential vocabulary sections")
        
        # Parse each section
        entries = []
        for i, section in enumerate(sections):
            entry = self.parse_vocabulary_section(section)
            if entry:
                entries.append(entry)
                print(f"✅ Extracted: {entry.word_or_phrase} (confidence: {entry.confidence_score:.2f})")
            else:
                print(f"⚠️ Skipped section {i+1} (couldn't extract vocabulary)")
        
        print(f"🎉 Successfully extracted {len(entries)} vocabulary entries")
        return entries
    
    def export_to_csv(self, entries: List[VocabularyEntry], output_path: Path):
        """Export entries to CSV format for import"""
        with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = [
                'word_or_phrase', 'definition', 'language', 'pronunciation',
                'etymology', 'memory_hint', 'difficulty_level', 'related_words'
            ]
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            
            for entry in entries:
                # Convert lists to comma-separated strings
                related_words = ', '.join(entry.related_words) if entry.related_words else ''
                
                # Combine English equivalents into memory hint
                memory_hint = entry.memory_hint or ''
                if entry.english_equivalents:
                    equiv_hint = f"English: {', '.join(entry.english_equivalents)}"
                    memory_hint = f"{memory_hint} {equiv_hint}".strip()
                
                writer.writerow({
                    'word_or_phrase': entry.word_or_phrase,
                    'definition': entry.definition,
                    'language': entry.language,
                    'pronunciation': entry.pronunciation or '',
                    'etymology': entry.etymology or '',
                    'memory_hint': memory_hint,
                    'difficulty_level': entry.difficulty_level,
                    'related_words': related_words
                })
    
    def export_to_json(self, entries: List[VocabularyEntry], output_path: Path):
        """Export entries to JSON format"""
        json_data = []
        for entry in entries:
            entry_dict = asdict(entry)
            # Remove None values and internal fields
            clean_dict = {k: v for k, v in entry_dict.items() 
                         if v is not None and k not in ['confidence_score', 'image_description']}
            json_data.append(clean_dict)
        
        with open(output_path, 'w', encoding='utf-8') as jsonfile:
            json.dump(json_data, jsonfile, indent=2, ensure_ascii=False)

def main():
    parser = argparse.ArgumentParser(description='Parse vocabulary from .docx files')
    parser.add_argument('input_file', help='Input .docx file path')
    parser.add_argument('-o', '--output', help='Output file path (CSV or JSON)')
    parser.add_argument('-f', '--format', choices=['csv', 'json'], default='csv',
                       help='Output format (default: csv)')
    parser.add_argument('--min-confidence', type=float, default=0.3,
                       help='Minimum confidence score for entries (default: 0.3)')
    
    args = parser.parse_args()
    
    input_path = Path(args.input_file)
    if not input_path.exists():
        print(f"❌ Error: File {input_path} not found")
        return
    
    # Parse document
    parser_instance = AdvancedVocabularyParser()
    entries = parser_instance.parse_docx(input_path)
    
    # Filter by confidence
    filtered_entries = [e for e in entries if e.confidence_score >= args.min_confidence]
    if len(filtered_entries) < len(entries):
        print(f"🔍 Filtered to {len(filtered_entries)} entries (min confidence: {args.min_confidence})")
    
    if not filtered_entries:
        print("❌ No vocabulary entries found with sufficient confidence")
        return
    
    # Determine output path
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.with_suffix(f'.{args.format}')
    
    # Export
    if args.format == 'csv':
        parser_instance.export_to_csv(filtered_entries, output_path)
    else:
        parser_instance.export_to_json(filtered_entries, output_path)
    
    print(f"✅ Exported {len(filtered_entries)} entries to {output_path}")
    
    # Show sample of extracted data
    print("\n📋 Sample entries:")
    for entry in filtered_entries[:3]:
        print(f"  • {entry.word_or_phrase}: {entry.definition[:100]}...")
        if entry.etymology:
            print(f"    Etymology: {entry.etymology[:80]}...")
        if entry.english_equivalents:
            print(f"    English: {', '.join(entry.english_equivalents)}")
        print()

if __name__ == '__main__':
    main()