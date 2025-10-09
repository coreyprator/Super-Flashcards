# backend/scripts/analyze_word_frequency.py
import re
from collections import Counter
from pathlib import Path
import csv
import argparse

try:
    from docx import Document
except ImportError:
    print("Please install python-docx: pip install python-docx")
    exit(1)

def analyze_docx(file_path, language="French"):
    """Extract word frequencies from .docx file"""
    print(f"Analyzing: {file_path}")
    
    doc = Document(file_path)
    all_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text.strip())
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    all_text.append(cell.text.strip())
    
    # Extract words
    words = []
    for segment in all_text:
        if len(segment) > 50:  # Skip long explanations
            continue
        if any(h in segment for h in ['Étymologie', 'Contexte', 'Équivalent']):
            continue
        
        matches = re.findall(r'\b[a-zA-ZàâäæçéèêëïîôùûüÿœÀÂÄÆÇÉÈÊËÏÎÔÙÛÜŸŒ\-\']+\b', segment)
        for word in matches:
            if len(word) >= 3 and not word.isupper():
                words.append(word.lower())
    
    frequency = Counter(words)
    print(f"Found {len(frequency)} unique words")
    
    return frequency

def export_csv(frequency, output_file, language="French", min_freq=1, max_words=None):
    """Export to CSV"""
    filtered = {w: c for w, c in frequency.items() if c >= min_freq}
    sorted_words = sorted(filtered.items(), key=lambda x: x[1], reverse=True)
    
    if max_words:
        sorted_words = sorted_words[:max_words]
    
    with open(output_file, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(['word', 'frequency', 'language', 'status'])
        for word, count in sorted_words:
            writer.writerow([word, count, language, 'pending'])
    
    print(f"Exported {len(sorted_words)} words to {output_file}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Analyze word frequency")
    parser.add_argument('input_file', help='Input .docx file')
    parser.add_argument('-o', '--output', default='word_frequency.csv')
    parser.add_argument('--language', default='French')
    parser.add_argument('--min-freq', type=int, default=1)
    parser.add_argument('--max-words', type=int, default=None)
    
    args = parser.parse_args()
    
    freq = analyze_docx(args.input_file, args.language)
    export_csv(freq, args.output, args.language, args.min_freq, args.max_words)
    
    print(f"\nTop 10 words:")
    for word, count in freq.most_common(10):
        print(f"  {word}: {count}")
