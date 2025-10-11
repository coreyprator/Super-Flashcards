"""
Document Parser API endpoint
Handles both .docx files and plain text parsing for vocabulary extraction
"""

from fastapi import APIRouter, UploadFile, File, HTTPException, Form, Depends
from sqlalchemy.orm import Session
from typing import List, Dict, Any, Optional
import tempfile
import shutil
from pathlib import Path
import json
import re
from dataclasses import dataclass, asdict

from ..database import get_db
from ..models import Language

router = APIRouter()

@dataclass
class VocabularyEntry:
    """Structured vocabulary entry"""
    word_or_phrase: str
    definition: str
    language: str = "french"
    pronunciation: Optional[str] = None
    etymology: Optional[str] = None
    memory_hint: Optional[str] = None
    difficulty_level: int = 1
    related_words: Optional[List[str]] = None
    english_equivalents: Optional[List[str]] = None
    expressions: Optional[List[str]] = None
    confidence_score: float = 1.0

class SmartVocabularyParser:
    """Smart parser for structured vocabulary content"""
    
    def extract_main_term(self, text: str) -> Optional[str]:
        """Extract the main vocabulary term"""
        lines = text.split('\n')
        for line in lines[:3]:  # Check first 3 lines
            line = line.strip()
            if not line:
                continue
            
            # Look for standalone terms
            if len(line.split()) <= 3 and len(line) > 2:
                # Clean punctuation
                term = re.sub(r'[«»""\(\)\[\]{}:]', '', line).strip()
                if term and not term.isupper():
                    return term.lower()
        
        return None
    
    def extract_definition(self, text: str, main_term: str) -> str:
        """Extract the definition text"""
        lines = text.split('\n')
        definition_lines = []
        
        # Skip lines until we find content (after the term)
        start_collecting = False
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Start collecting after we see the main term or skip initial headers
            if main_term.lower() in line.lower() or start_collecting:
                start_collecting = True
                # Stop at etymology or equivalents sections
                if any(keyword in line.lower() for keyword in ['étymologie', 'équivalent', 'mots apparentés']):
                    break
                if line and not line.startswith('●'):
                    definition_lines.append(line)
        
        definition = ' '.join(definition_lines).strip()
        # Remove the term itself from the definition
        definition = re.sub(f'^{re.escape(main_term)}[^a-zA-Z]*', '', definition, flags=re.IGNORECASE)
        return definition[:500] if definition else f"Terme français: {main_term}"
    
    def extract_etymology(self, text: str) -> Optional[str]:
        """Extract etymology information"""
        etymology_section = re.search(
            r'Étymologie\s*:?\s*(.+?)(?=\n\n|Équivalent|Mots apparentés|$)', 
            text, re.DOTALL | re.IGNORECASE
        )
        
        if etymology_section:
            etymology = etymology_section.group(1).strip()
            # Clean up bullet points
            etymology = re.sub(r'●\s*', '• ', etymology)
            etymology = re.sub(r'\s+', ' ', etymology)
            return etymology[:300]
        
        return None
    
    def extract_english_equivalents(self, text: str) -> List[str]:
        """Extract English equivalent words"""
        equivalents = []
        
        # Look for English equivalents sections
        patterns = [
            r'Équivalent[^:]*:?\s*["\"]([^"\"]+)["\"]',
            r'en anglais[^:]*:?\s*["\"]([^"\"]+)["\"]',
            r'●\s*Équivalent\s*:\s*["\"]([^"\"]+)["\"]'
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                equiv_text = match.group(1).strip()
                # Split on 'ou' or 'or'
                words = re.split(r'\s+ou\s+|\s+or\s+|,\s*', equiv_text)
                for word in words:
                    clean_word = re.sub(r'["\(\)\.]+', '', word.strip())
                    if clean_word and len(clean_word) > 1:
                        equivalents.append(clean_word)
        
        return equivalents[:5]
    
    def extract_related_words(self, text: str) -> List[str]:
        """Extract French related words"""
        related = []
        
        # Look for French related words section
        french_section = re.search(
            r'Mots apparentés en français[^:]*:?\s*(.+?)(?=\n\n|$)', 
            text, re.DOTALL | re.IGNORECASE
        )
        
        if french_section:
            related_text = french_section.group(1)
            # Extract bullet point items
            words = re.findall(r'●\s*([^\n●]+)', related_text)
            for word in words:
                clean_word = word.strip()
                if clean_word and len(clean_word) > 1:
                    related.append(clean_word)
        
        return related[:8]
    
    def determine_difficulty(self, entry: VocabularyEntry) -> int:
        """Estimate difficulty level"""
        score = 1
        
        if entry.etymology and len(entry.etymology) > 100:
            score += 1
        if entry.related_words and len(entry.related_words) > 3:
            score += 1
        if len(entry.definition) > 150:
            score += 1
        if entry.english_equivalents and len(entry.english_equivalents) > 2:
            score += 1
        
        return min(score, 5)
    
    def parse_text_sections(self, text: str) -> List[VocabularyEntry]:
        """Parse text into vocabulary entries"""
        # Split on separator lines or double newlines
        sections = re.split(r'_{8,}|\n\s*\n\s*\n', text)
        
        entries = []
        for section in sections:
            section = section.strip()
            if len(section) < 50:  # Skip short sections
                continue
            
            # Extract main term
            main_term = self.extract_main_term(section)
            if not main_term:
                continue
            
            # Create entry
            entry = VocabularyEntry(
                word_or_phrase=main_term,
                definition=self.extract_definition(section, main_term),
                language="french"
            )
            
            # Extract additional info
            entry.etymology = self.extract_etymology(section)
            entry.english_equivalents = self.extract_english_equivalents(section)
            entry.related_words = self.extract_related_words(section)
            entry.difficulty_level = self.determine_difficulty(entry)
            
            # Create memory hint from English equivalents
            if entry.english_equivalents:
                entry.memory_hint = f"English: {', '.join(entry.english_equivalents)}"
            
            # Calculate confidence
            confidence = 0.6  # Base confidence
            if entry.etymology:
                confidence += 0.2
            if entry.english_equivalents:
                confidence += 0.1
            if entry.related_words:
                confidence += 0.1
            
            entry.confidence_score = min(confidence, 1.0)
            entries.append(entry)
        
        return entries

@router.post("/parse")
async def parse_document(
    file: UploadFile = File(...),
    language: str = Form("french"),
    min_confidence: float = Form(0.5),
    db: Session = Depends(get_db)
):
    """
    Parse a document (text or .docx) and extract vocabulary entries.
    Returns structured vocabulary data ready for import.
    """
    
    # Validate file
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    file_extension = file.filename.lower().split('.')[-1]
    if file_extension not in ['txt', 'docx']:
        raise HTTPException(status_code=400, detail="File must be .txt or .docx format")
    
    # Read file content
    try:
        content = await file.read()
        
        if file_extension == 'docx':
            # Handle .docx files
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(content)
                tmp_file.flush()
                
                try:
                    from docx import Document
                    doc = Document(tmp_file.name)
                    
                    # Extract text from docx
                    text_content = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_content.append(para.text)
                    
                    # Extract table content
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text_content.append(cell.text)
                    
                    file_content = '\n'.join(text_content)
                    
                finally:
                    Path(tmp_file.name).unlink()  # Clean up temp file
        else:
            # Handle text files
            file_content = content.decode('utf-8')
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading file: {str(e)}")
    
    # Parse content
    parser = SmartVocabularyParser()
    entries = parser.parse_text_sections(file_content)
    
    # Filter by confidence
    filtered_entries = [e for e in entries if e.confidence_score >= min_confidence]
    
    # Get available languages for validation
    languages = db.query(Language).all()
    language_names = {lang.name.lower(): lang.name for lang in languages}
    
    # Prepare response
    result = {
        "success": True,
        "message": f"Extracted {len(filtered_entries)} vocabulary entries from {len(entries)} potential matches",
        "total_sections": len(entries),
        "extracted_entries": len(filtered_entries),
        "min_confidence": min_confidence,
        "available_languages": list(language_names.values()),
        "entries": []
    }
    
    # Convert entries to dictionaries for JSON response
    for entry in filtered_entries:
        entry_dict = {
            "word_or_phrase": entry.word_or_phrase,
            "definition": entry.definition,
            "language": language_names.get(entry.language.lower(), entry.language),
            "pronunciation": entry.pronunciation,
            "etymology": entry.etymology,
            "memory_hint": entry.memory_hint,
            "difficulty_level": entry.difficulty_level,
            "related_words": entry.related_words,
            "confidence_score": round(entry.confidence_score, 2)
        }
        result["entries"].append(entry_dict)
    
    return result

@router.get("/sample")
async def get_sample_format():
    """Get a sample of the expected document format"""
    sample = """étatique
Le terme "étatique" provient du vocabulaire administratif et politique français et est utilisé pour décrire ce qui est relatif à l'État.

Étymologie :
●	Origine latine : Le mot "État" dérive du latin status, signifiant "position, condition".

Équivalent ou mots apparentés en anglais :
●	Équivalent : "State-related" ou "statist".

Mots apparentés en français :
●	Administration
●	Gouvernemental
●	Public

________________________________________

inadaptées  
Le terme "inadaptées" est un adjectif qui signifie "qui n'est pas approprié à une situation."

Étymologie :
●	Origine latine : Composé du préfixe "in-" et du verbe latin adaptare.

Équivalent ou mots apparentés en anglais :
●	Équivalent : "Unsuitable," "inappropriate."

Mots apparentés en français :
●	Inappropriées
●	Inadéquates"""
    
    return {
        "sample_format": sample,
        "instructions": [
            "Each vocabulary entry should be separated by underscores or double line breaks",
            "Start with the main term on its own line",
            "Include definition in the first paragraph",
            "Add 'Étymologie :' section with bullet points (●)",
            "Add 'Équivalent ou mots apparentés en anglais :' section",
            "Add 'Mots apparentés en français :' section",
            "Use bullet points (●) for lists"
        ]
    }